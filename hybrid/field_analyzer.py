# # hybrid/field_analyzer.py
# import fitz  # PyMuPDF
# import os
# from pdf2image import convert_from_path
# from PIL import Image  # Importamos Pillow para combinar las páginas

# def extract_text_from_pdf(pdf_path):
#     """Extrae el texto plano del PDF de forma tradicional."""
#     try:
#         doc = fitz.open(pdf_path)
#         text = ""
#         for page in doc:
#             text += page.get_text()
#         return text
#     except Exception as e:
#         print(f" Error al leer texto del PDF: {e}")
#         return ""

# def convert_pdf_first_page_to_image(pdf_path, output_image_path="documents/temp_page.png"):
#     """
#     Convierte TODAS las páginas del PDF en imágenes PNG y las une 
#     verticalmente en una sola imagen larga para el análisis de la IA.
#     """
#     try:
#         ruta_poppler = r"C:\Users\GSF\AppData\Local\Microsoft\WinGet\Packages\oschwartz10612.Poppler_Microsoft.Winget.Source_8wekyb3d8bbwe\poppler-25.07.0\Library\bin"
        
#         #  Quitamos 'first_page' y 'last_page' para que convierta todo el documento
#         print(" Extrayendo visualmente todas las páginas del PDF...")
#         images = convert_from_path(pdf_path, poppler_path=ruta_poppler)
        
#         if images:
#             # Calcular el ancho máximo y la altura total combinada
#             max_width = max(img.width for img in images)
#             total_height = sum(img.height for img in images)
            
#             # Crear un lienzo en blanco del tamaño de todas las páginas juntas
#             combined_image = Image.new('RGB', (max_width, total_height))
            
#             # Pegar una página debajo de la otra
#             current_y = 0
#             for img in images:
#                 combined_image.paste(img, (0, current_y))
#                 current_y += img.height
            
#             # Asegurar que la carpeta exista y guardar la súper imagen
#             os.makedirs(os.path.dirname(output_image_path), exist_ok=True)
#             combined_image.save(output_image_path, 'PNG')
            
#             print(f" Súper imagen generada (Páginas combinadas: {len(images)}) en: {output_image_path}")
#             return output_image_path
            
#     except Exception as e:
#         print(f" Error al convertir el PDF completo a imagen: {e}")
#     return None


# # hybrid/field_analyzer.py
# import fitz  # PyMuPDF

# def extract_text_with_page_mapping(pdf_path):
#     """
#     Extrae el texto completo y crea un mapa de palabras vinculadas 
#     a su número de página real para poder rastrear discrepancias.
#     """
#     texto_completo = ""
#     mapa_paginas = [] # Lista de tuplas (palabra, pagina)
    
#     try:
#         doc = fitz.open(pdf_path)
#         for num_pag, page in enumerate(doc, start=1):
#             texto_pagina = page.get_text("text")
#             texto_completo += texto_pagina + "\n"
            
#             # Limpieza básica para el mapeo por palabras
#             palabras = texto_pagina.split()
#             for palabra in palabras:
#                 # Almacenamos la palabra limpia y la página a la que pertenece
#                 palabra_limpia = palabra.strip().replace("\\", "")
#                 if palabra_limpia:
#                     mapa_paginas.append((palabra_limpia, num_pag))
#         doc.close()
#     except Exception as e:
#         print(f"Error al mapear texto del PDF: {e}")
        
#     return texto_completo, mapa_paginas

# def extract_text_from_pdf(pdf_path):
#     """Mantiene compatibilidad con llamadas simples de texto plano."""
#     texto, _ = extract_text_with_page_mapping(pdf_path)
#     return texto

# # hybrid/field_analyzer.py
# import fitz  # PyMuPDF
# import re

# def limpiar_linea(texto: str) -> str:
#     """Limpia el ruido de una línea individual sin perder su flujo."""
#     if not texto:
#         return ""
#     texto = texto.replace("\\", "")
#     texto = texto.replace("\r", " ").replace("\n", " ").replace("\t", " ")
#     return re.sub(r"\s+", " ", texto).strip()

# def extract_text_with_page_mapping(pdf_path):
#     """
#     Extrae el texto organizándolo por líneas limpias, mapeando 
#     cada línea a su número de página real.
#     """
#     texto_completo = ""
#     lineas_mapeadas = []  # Lista de tuplas (linea_limpia, num_pagina)
    
#     try:
#         doc = fitz.open(pdf_path)
#         for num_pag, page in enumerate(doc, start=1):
#             texto_pagina = page.get_text("text")
#             texto_completo += texto_pagina + "\n"
            
#             # Procesar por líneas físicas del PDF
#             for line in texto_pagina.split('\n'):
#                 linea_limpia = limpiar_linea(line)
#                 # Omitir headers/footers repetitivos y líneas vacías
#                 if not linea_limpia or "Bläckfisk" in linea_limpia or "Estructuras Financieras" in linea_limpia:
#                     continue
#                 lineas_mapeadas.append((linea_limpia, num_pag))
#         doc.close()
#     except Exception as e:
#         print(f" Error al mapear texto del PDF: {e}")
        
#     return texto_completo, lineas_mapeadas

# # hybrid/field_analyzer.py
# import fitz  # PyMuPDF
# import re
# import os
# from docx import Document

# def limpiar_linea(texto: str) -> str:
#     """Normaliza por completo el texto eliminando ruidos visuales."""
#     if not texto:
#         return ""
#     texto = texto.replace("\\", "")
#     texto = texto.replace("\r", " ").replace("\n", " ").replace("\t", " ")
#     return re.sub(r"\s+", " ", texto).strip()

# def _extract_from_pdf(pdf_path):
#     """Extractor interno para archivos PDF."""
#     texto_completo = ""
#     lineas_mapeadas = []
    
#     doc = fitz.open(pdf_path)
#     for num_pag, page in enumerate(doc, start=1):
#         texto_pagina = page.get_text("text")
#         texto_completo += texto_pagina + "\n"
        
#         for line in texto_pagina.split('\n'):
#             linea_limpia = limpiar_linea(line)
#             if not linea_limpia:
#                 continue
#             lineas_mapeadas.append((linea_limpia, num_pag))
#     doc.close()
#     return texto_completo, lineas_mapeadas

# def _extract_from_docx(docx_path):
#     """Extractor interno para archivos DOCX (Word)."""
#     texto_completo = ""
#     lineas_mapeadas = []
    
#     doc = Document(docx_path)
#     # En Word usamos el índice del párrafo como sustituto del número de página
#     for idx_parrafo, paragraph in enumerate(doc.paragraphs, start=1):
#         texto_parrafo = paragraph.text
#         if not texto_parrafo.strip():
#             continue
            
#         texto_completo += texto_parrafo + "\n"
        
#         # Dividir por si hay saltos de línea internos dentro del mismo párrafo
#         for line in texto_parrafo.split('\n'):
#             linea_limpia = limpiar_linea(line)
#             if not linea_limpia:
#                 continue
#             # Guardamos el texto y una etiqueta clara de ubicación
#             lineas_mapeadas.append((linea_limpia, f"Párrafo {idx_parrafo}"))
            
#     return texto_completo, lineas_mapeadas

# def extract_text_with_page_mapping(file_path):
#     """
#     EXTRACTOR UNIVERSAL: Detecta automáticamente el formato del archivo 
#     y extrae sus líneas de forma compatible para el comparador.
#     """
#     ext = os.path.splitext(file_path)[1].lower()
    
#     if ext == '.pdf':
#         return _extract_from_pdf(file_path)
#     elif ext == '.docx':
#         return _extract_from_docx(file_path)
#     else:
#         raise ValueError(f"❌ Formato de archivo no soportado de momento: {ext}")


# hybrid/field_analyzer.py
import fitz  # PyMuPDF
import re
import os
from docx import Document

def limpiar_linea(texto: str) -> str:
    """Normaliza por completo el texto eliminando ruidos visuales."""
    if not texto:
        return ""
    texto = texto.replace("\\", "")
    texto = texto.replace("\r", " ").replace("\n", " ").replace("\t", " ")
    return re.sub(r"\s+", " ", texto).strip()

def es_cabecera_pura(texto: str) -> bool:
    """
    Detecta si una línea aislada del PDF contiene únicamente datos del formulario
    para evitar que se concatene de forma caótica en los bloques contractuales.
    """
    t = texto.lower()
    # Si la línea es solo etiquetas de formulario y números/códigos
    if "reca:" in t or "número de cliente" in t or "número de contrato" in t or "número de referencia" in t:
        # Si no tiene palabras de contrato largas, es una cabecera pura
        if not any(palabra in t for palabra in ["promete", "incondicionalmente", "pagar", "suscrito"]):
            return True
    return False

def _extract_from_pdf(pdf_path):
    """Extractor interno para archivos PDF con limpieza de ruido en cabeceras."""
    texto_completo = ""
    lineas_mapeadas = []
    
    doc = fitz.open(pdf_path)
    for num_pag, page in enumerate(doc, start=1):
        texto_pagina = page.get_text("text")
        
        for line in texto_pagina.split('\n'):
            linea_limpia = limpiar_linea(line)
            if not linea_limpia:
                continue
                
            # Si es una línea de datos dinámicos pura que el extractor mezcló, 
            # la separamos del flujo principal convirtiéndola en un marcador estándar
            if es_cabecera_pura(linea_limpia):
                lineas_mapeadas.append(("[DATOS_CABECERA_OMITIDOS]", num_pag))
                texto_completo += "[DATOS_CABECERA_OMITIDOS]\n"
                continue
                
            lineas_mapeadas.append((linea_limpia, num_pag))
            texto_completo += linea_limpia + "\n"
            
    doc.close()
    return texto_completo, lineas_mapeadas

def _extract_from_docx(docx_path):
    """Extractor interno para archivos DOCX (Word)."""
    texto_completo = ""
    lineas_mapeadas = []
    
    doc = Document(docx_path)
    for idx_parrafo, paragraph in enumerate(doc.paragraphs, start=1):
        texto_parrafo = paragraph.text
        if not texto_parrafo.strip():
            continue
            
        texto_completo += texto_parrafo + "\n"
        
        for line in texto_parrafo.split('\n'):
            linea_limpia = limpiar_linea(line)
            if not linea_limpia:
                continue
            lineas_mapeadas.append((linea_limpia, f"Párrafo {idx_parrafo}"))
            
    return texto_completo, lineas_mapeadas

def extract_text_with_page_mapping(file_path):
    """EXTRACTOR UNIVERSAL: Detecta automáticamente el formato del archivo."""
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext == '.pdf':
        return _extract_from_pdf(file_path)
    elif ext == '.docx':
        return _extract_from_docx(file_path)
    else:
        raise ValueError(f"Formato de archivo no soportado: {ext}")