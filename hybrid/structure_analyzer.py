# # hybrid/structure_analyzer.py

# import re

# # Mi plantilla maestra basada en títulos
# DOCUMENT_TEMPLATES = {
#     "ANEXO DE DISPOSICIONES LEGALES": [
#         r"CÓDIGO CIVIL FEDERAL",
#         r"DOCUMENTOS:",
#         r"LEY DE INSTITUCIONES DE CRÉDITO",
#         r"LEY DE PROTECCIÓN Y DEFENSA AL USUARIO DE SERVICIOS FINANCIEROS\.",
#         r"LEY GENERAL DE ORGANIZACIONES Y ACTIVIDADES AUXILIARES DEL CRÉDITO",
#         r"LEY GENERAL DE TÍTULOS Y OPERACIONES DE CRÉDITO",
#         r"LEY PARA LA TRANSPARENCIA Y ORDENAMIENTO DE LOS SERVICIOS FINANCIEROS",
#         r"LEY PARA REGULAR LAS SOCIEDADES DE INFORMACIÓN CREDITICIA",
#         r"RECA:\s*.*"  # Regex para aceptar cualquier RECA variable
#     ]
# }

# def identify_document_by_title(text):
#     """Busca si el título del documento coincide con alguna plantilla."""
#     for title in DOCUMENT_TEMPLATES.keys():
#         if title in text:
#             return title
#     return None

# def analyze_document_structure(actual_text, doc_title):
#     """Compara las secciones del texto real contra la plantilla esperada."""
#     if not doc_title or doc_title not in DOCUMENT_TEMPLATES:
#         return {"error": f"Plantilla no encontrada para el título: {doc_title}"}
    
#     regex_esperadas = DOCUMENT_TEMPLATES[doc_title]
    
#     # Formateamos visualmente las secciones esperadas limpiando los caracteres regex
#     secciones_esperadas = [
#         "RECA: [ ]" if r"RECA:" in rx else rx.replace("\\.", ".") 
#         for rx in regex_esperadas
#     ]
    
#     secciones_actuales = []
#     secciones_comunes = []
#     secciones_perdidas = []
    
#     for i, regex in enumerate(regex_esperadas):
#         match = re.search(regex, actual_text)
#         nombre_legible = secciones_esperadas[i]
        
#         if match:
#             texto_real = match.group(0).strip()
#             secciones_actuales.append(texto_real)
            
#             # Si es el RECA dinámico, guardamos el formato limpio en comunes
#             if r"RECA:" in regex:
#                 secciones_comunes.append("RECA: [ ]")
#             else:
#                 secciones_comunes.append(texto_real)
#         else:
#             secciones_perdidas.append(nombre_legible)
            
#     # Calcular la puntuación estructural para reporte de QA
#     total = len(secciones_esperadas)
#     encontradas = len(secciones_comunes)
#     puntuacion = (encontradas / total) * 100 if total > 0 else 0
    
#     return {
#         "sections_esperadas": secciones_esperadas,
#         "secciones_actuales": secciones_actuales,
#         "secciones_comunes": secciones_comunes,
#         "secciones_perdidas": secciones_perdidas,
#         "puntuacion_estructura": round(puntuacion, 2)
#     }

# ## CALCULARA LOS BLOQUES DE TEXTO DE CADA SECCION

# # hybrid/structure_analyzer.py
# import re

# DOCUMENT_TEMPLATES = {
#     "ANEXO DE DISPOSICIONES LEGALES": [
#         r"CÓDIGO CIVIL FEDERAL",
#         r"DOCUMENTOS:",
#         r"LEY DE INSTITUCIONES DE CRÉDITO",
#         r"LEY DE PROTECCIÓN Y DEFENSA AL USUARIO DE SERVICIOS FINANCIEROS\.",
#         r"LEY GENERAL DE ORGANIZACIONES Y ACTIVIDADES AUXILIARES DEL CRÉDITO",
#         r"LEY GENERAL DE TÍTULOS Y OPERACIONES DE CRÉDITO",
#         r"LEY PARA LA TRANSPARENCIA Y ORDENAMIENTO DE LOS SERVICIOS FINANCIEROS",
#         r"LEY PARA REGULAR LAS SOCIEDADES DE INFORMACIÓN CREDITICIA",
#     ]
# }

# def identify_document_by_title(text):
#     for title in DOCUMENT_TEMPLATES.keys():
#         if title in text:
#             return title
#     return None

# def analyze_document_structure(actual_text, doc_title):
#     if not doc_title or doc_title not in DOCUMENT_TEMPLATES:
#         return {"error": f"Plantilla no encontrada para el título: {doc_title}"}
    
#     regex_esperadas = DOCUMENT_TEMPLATES[doc_title]
    
#     # Formateo de nombres limpios para las claves de salida
#     secciones_esperadas = [
#         "RECA: [ ]" if r"RECA:" in rx else rx.replace("\\.", ".") 
#         for rx in regex_esperadas
#     ]
    
#     secciones_encontradas = []
#     secciones_faltantes = []
#     # Diccionario donde guardaremos {"Nombre de la Sección": "Texto interno extraído"}
#     contenido_secciones = {}
    
#     # Encontrar las posiciones de inicio de cada sección en el documento real
#     matches = []
#     for i, regex in enumerate(regex_esperadas):
#         match = re.search(regex, actual_text)
#         nombre_legible = secciones_esperadas[i]
        
#         if match:
#             matches.append({
#                 "nombre": nombre_legible,
#                 "inicio": match.start(),
#                 "fin": match.end(),
#                 "texto_titulo": match.group(0).strip()
#             })
#             secciones_encontradas.append(texto_titulo_real := match.group(0).strip())
#         else:
#             secciones_faltantes.append(nombre_legible)

#     # Ordenar los matches por su aparición en el texto para poder recortar entre ellos
#     matches = sorted(matches, key=lambda x: x["inicio"])
    
#     # Extraer el contenido de fondo que pertenece a cada sección
#     for index, match_actual in enumerate(matches):
#         inicio_contenido = match_actual["fin"]
        
#         # Si hay una sección siguiente, el contenido termina donde empieza la otra.
#         # Si es la última sección, el contenido termina al final de todo el texto del PDF.
#         if index + 1 < len(matches):
#             fin_contenido = matches[index + 1]["inicio"]
#             texto_interno = actual_text[inicio_contenido:fin_contenido].strip()
#         else:
#             texto_interno = actual_text[inicio_contenido:].strip()
            
#         contenido_secciones[match_actual["nombre"]] = texto_interno

#     total = len(secciones_esperadas)
#     encontradas = len(secciones_encontradas)
#     puntuacion = (encontradas / total) * 100 if total > 0 else 0
    
#     return {
#         "secciones_esperadas": secciones_esperadas,
#         "secciones_encontradas": secciones_encontradas,
#         "secciones_faltantes": secciones_faltantes,
#         "puntuacion_estructura": round(puntuacion, 2),
#         "contenido_secciones": contenido_secciones  # <-- Datos listos para comparar contenido
#     }

# # hybrid/structure_analyzer.py
# import re
# import ollama
# import json

# def discover_sections_with_ai(document_text, model_name="deepseek-r1:latest"):
#     """
#     Lee el texto de un documento esperado y le pide a la IA que descubra 
#     cuáles son los títulos de las secciones, leyes o apartados principales.
#     """
#     print(" Descubriendo secciones del documento original con IA...")
    
#     prompt = f"""
#     Eres un extractor de metadatos de alta precisión para documentos legales y financieros.
#     Tu única tarea es identificar y listar TODOS los títulos de secciones, nombres de leyes, capítulos o identificadores únicos (como códigos RECA) presentes en el texto.

#     REGLAS ESTRICTAS:
#     1. Extrae el título EXACTAMENTE como aparece escrito (mismas palabras).
#     2. No ignores leyes secundarias o menciones de códigos regulatorios.
#     3. Si detectas titulos con "[]" o similares, no los tomes encuenta (ej. "RECA: J3B94U32" o "RECA: [ ]").
#     4. Ignora párrafos de contenido común, solo extrae los encabezados/títulos.

#     EJEMPLO DE SALIDA ESPERADA:
#     {{
#         "secciones_detectadas": [
#             "ANEXO DE DISPOSICIONES LEGALES",
#             "CÓDIGO CIVIL FEDERAL",
#             "DOCUMENTOS:",
#             "LEY DE INSTITUCIONES DE CRÉDITO",
#             "RECA: J3B94U32"
#         ]
#     }}

#     TEXTO DEL DOCUMENTO A ANALIZAR:
#     {document_text[:7000]}

#     Responde SÓLO el JSON plano, sin formato markdown (NO uses ```json):
#     """
    
#     # Dentro de discover_sections_with_ai o compare_section_content_with_ai:
#     try:
#         response = ollama.generate(
#             model=model_name, 
#             prompt=prompt,
#             options={
#                 "temperature": 0.0,  # <-- Fuerza al modelo a ser ultra preciso y no inventar/omitir nada
#                 "top_p": 0.1
#             }
#         )
        
#         clean_response = response['response'].strip().replace("```json", "").replace("```", "")
#         data = json.loads(clean_response)
#         return data.get("secciones_detectadas", [])
#     except Exception as e:
#         print(f"Error al descubrir secciones con IA: {e}")
#         return []

# def analyze_document_structure_dynamic(actual_text, secciones_esperadas):
#     """
#     Compara las secciones del texto real contra la lista dinámica que la IA descubrió.
#     """
#     secciones_encontradas = []
#     secciones_faltantes = []
#     contenido_secciones = {}
#     matches = []
    
#     for nombre_seccion in secciones_esperadas:
#         # Si la sección contiene algo variable como RECA, creamos un regex flexible
#         if "RECA" in nombre_seccion.upper():
#             regex = r"RECA:\s*.*"
#         else:
#             # Escapamos caracteres especiales del título por seguridad
#             regex = re.escape(nombre_seccion)
            
#         match = re.search(regex, actual_text, re.IGNORECASE)
        
#         if match:
#             texto_real_encontrado = match.group(0).strip()
#             matches.append({
#                 "nombre": nombre_seccion,
#                 "inicio": match.start(),
#                 "fin": match.end(),
#                 "texto_titulo": texto_real_encontrado
#             })
#             secciones_encontradas.append(texto_real_encontrado)
#         else:
#             secciones_faltantes.append(nombre_seccion)

#     # Ordenar por aparición para poder recortar los contenidos intermedios
#     matches = sorted(matches, key=lambda x: x["inicio"])
    
#     for index, match_actual in enumerate(matches):
#         inicio_contenido = match_actual["fin"]
        
#         if index + 1 < len(matches):
#             fin_contenido = matches[index + 1]["inicio"]
#             texto_interno = actual_text[inicio_contenido:fin_contenido].strip()
#         else:
#             texto_interno = actual_text[inicio_contenido:].strip()
            
#         contenido_secciones[match_actual["nombre"]] = texto_interno

#     total = len(secciones_esperadas)
#     encontradas = len(secciones_encontradas)
#     puntuacion = (encontradas / total) * 100 if total > 0 else 0
    
#     return {
#         "secciones_esperadas": secciones_esperadas,
#         "secciones_encontradas": secciones_encontradas,
#         "secciones_faltantes": secciones_faltantes,
#         "puntuacion_estructura": round(puntuacion, 2),
#         "contenido_secciones": contenido_secciones
#     }


# # hybrid/structure_analyzer.py
# import re
# import ollama
# import json

# def discover_sections_with_ai(document_text, model_name="deepseek-r1:7b"):
#     print("\n Descubriendo secciones del documento original con IA...")
#     print(" Esperando respuesta de Ollama (Pensando...): \n")
    
#     prompt = f"""
#     Eres un extractor de metadatos de alta precisión para documentos legales y financieros.
#     Tu única tarea es identificar y listar TODOS los títulos de secciones, nombres de leyes, capítulos o identificadores únicos (como códigos RECA) presentes en el texto.

#     REGLAS ESTRICTAS:
#     1. Extrae el título EXACTAMENTE como aparece escrito (mismas palabras).
#     2. No ignores leyes secundarias o menciones de códigos regulatorios.
#     3. Si detectas un código RECA, inclúyelo completo (ej. "RECA: J3B94U32" o "RECA: [ ]").
#     4. Ignora párrafos de contenido común, solo extrae los encabezados/títulos.

#     TEXTO DEL DOCUMENTO A ANALIZAR:
#     {document_text[:4000]}

#     Responde SÓLO un JSON plano con la clave 'secciones_detectadas'. No uses formato markdown.
#     """
    
#     try:
#         response_stream = ollama.generate(
#             model=model_name, 
#             prompt=prompt, 
#             stream=True,
#             options={"temperature": 0.0}
#         )
        
#         full_response = ""
#         for chunk in response_stream:
#             texto_fragmento = chunk['response']
#             full_response += texto_fragmento
#             print(texto_fragmento, end="", flush=True)
            
#         print("\n\n Fin de la transmisión de la IA. Procesando JSON...")
        
#         if "</think>" in full_response:
#             full_response = full_response.split("</think>")[-1].strip()
            
#         clean_response = full_response.replace("```json", "").replace("```", "").strip()
#         data = json.loads(clean_response)
#         return data.get("secciones_detectadas", [])
        
#     except Exception as e:
#         print(f"\n Error al descubrir secciones con IA: {e}")
#         return []

# def analyze_document_structure_dynamic(actual_text, secciones_esperadas):
#     """
#     Compara las secciones del texto real contra la lista dinámica que la IA descubrió.
#     """
#     secciones_encontradas = []
#     secciones_faltantes = []
#     contenido_secciones = {}
#     matches = []
    
#     for nombre_seccion in secciones_esperadas:
#         if "RECA" in nombre_seccion.upper():
#             regex = r"RECA:\s*.*"
#         else:
#             regex = re.escape(nombre_seccion)
            
#         match = re.search(regex, actual_text, re.IGNORECASE)
        
#         if match:
#             texto_real_encontrado = match.group(0).strip()
#             matches.append({
#                 "nombre": nombre_seccion,
#                 "inicio": match.start(),
#                 "fin": match.end(),
#                 "texto_titulo": texto_real_encontrado
#             })
#             secciones_encontradas.append(texto_real_encontrado)
#         else:
#             secciones_faltantes.append(nombre_seccion)

#     matches = sorted(matches, key=lambda x: x["inicio"])
    
#     for index, match_actual in enumerate(matches):
#         inicio_contenido = match_actual["fin"]
        
#         if index + 1 < len(matches):
#             fin_contenido = matches[index + 1]["inicio"]
#             texto_interno = actual_text[inicio_contenido:fin_contenido].strip()
#         else:
#             texto_interno = actual_text[inicio_contenido:].strip()
            
#         contenido_secciones[match_actual["nombre"]] = texto_interno

#     total = len(secciones_esperadas)
#     encontradas = len(secciones_encontradas)
#     puntuacion = (encontradas / total) * 100 if total > 0 else 0
    
#     return {
#         "secciones_esperadas": secciones_esperadas,
#         "secciones_encontradas": secciones_encontradas,
#         "secciones_faltantes": secciones_faltantes,
#         "puntuacion_estructura": round(puntuacion, 2),
#         "contenido_secciones": contenido_secciones
#     }


# # hybrid/structure_analyzer.py
# import re
# import ollama
# import json

# # hybrid/structure_analyzer.py
# import re
# import ollama
# import json

# def discover_sections_with_vision_ai(image_path, model_name="minicpm-v"):
#     print(f"\n Analizando visualmente el documento con {model_name}...")
#     print(" Procesando imagen completa (Esto puede tardar un poco más por el tamaño)... \n")
    
#     prompt = """
#     Analiza la imagen de este documento legal. Identifica y extrae TODOS los títulos principales, encabezados en negrita y centradas al igual que nombres de leyes que estructuran el documento (por ejemplo: 'ANEXO DE DISPOSICIONES LEGALES', 'CÓDIGO CIVIL FEDERAL', 'LEY DE INSTITUCIONES DE CRÉDITO', o 'RECA: ...').

#     Responde EXCLUSIVAMENTE en un formato JSON válido con la clave 'secciones_detectadas'.
#     Ejemplo de formato requerido:
#     {
#         "secciones_detectadas": ["TITULO 1", "TITULO 2"]
#     }
#     """
    
#     try:
#         response = ollama.generate(
#             model=model_name,
#             prompt=prompt,
#             images=[image_path],
#             options={"temperature": 0.0}
#         )
        
#         full_response = response['response'].strip()
        
#         #  DEPURACIÓN: Esto te dejará ver EXACTAMENTE qué escribió la IA en la consola
#         print("================= RESPUESTA CRUDA DE LA IA =================")
#         print(full_response)
#         print("============================================================\n")
        
#         #  ESTRATEGIA DE EXTRACCIÓN ROBUSTA CON REGEX
#         # Buscamos cualquier cosa que esté entre las llaves { ... } incluyendo saltos de línea
#         json_match = re.search(r'\{.*\}', full_response, re.DOTALL)
        
#         if json_match:
#             json_puro = json_match.group(0)
#             data = json.loads(json_puro)
#             return data.get("secciones_detectadas", [])
#         else:
#             # Si no encontró llaves, intentamos limpiar formatos de texto comunes por si acaso
#             clean_response = full_response.replace("```json", "").replace("```", "").strip()
#             data = json.loads(clean_response)
#             return data.get("secciones_detectadas", [])
            
#     except Exception as e:
#         print(f" Error al procesar o parsear la respuesta de la IA: {e}")
#         return []

# # La función analyze_document_structure_dynamic se queda exactamente igual abajo...

# def analyze_document_structure_dynamic(actual_text, secciones_esperadas):
#     """Mantiene la misma lógica matemática que ya tenías para separar los textos"""
#     secciones_encontradas = []
#     secciones_faltantes = []
#     contenido_secciones = {}
#     matches = []
    
#     for nombre_seccion in secciones_esperadas:
#         if "RECA" in nombre_seccion.upper():
#             regex = r"RECA:\s*.*"
#         else:
#             regex = re.escape(nombre_seccion)
            
#         match = re.search(regex, actual_text, re.IGNORECASE)
        
#         if match:
#             texto_real_encontrado = match.group(0).strip()
#             matches.append({
#                 "nombre": nombre_seccion,
#                 "inicio": match.start(),
#                 "fin": match.end(),
#                 "texto_titulo": texto_real_encontrado
#             })
#             secciones_encontradas.append(texto_real_encontrado)
#         else:
#             secciones_faltantes.append(nombre_seccion)

#     matches = sorted(matches, key=lambda x: x["inicio"])
    
#     for index, match_actual in enumerate(matches):
#         inicio_contenido = match_actual["fin"]
        
#         if index + 1 < len(matches):
#             fin_contenido = matches[index + 1]["inicio"]
#             texto_interno = actual_text[inicio_contenido:fin_contenido].strip()
#         else:
#             texto_interno = actual_text[inicio_contenido:].strip()
            
#         contenido_secciones[match_actual["nombre"]] = texto_interno

#     total = len(secciones_esperadas)
#     encontradas = len(secciones_encontradas)
#     puntuacion = (encontradas / total) * 100 if total > 0 else 0
    
#     return {
#         "secciones_esperadas": secciones_esperadas,
#         "secciones_encontradas": secciones_encontradas,
#         "secciones_faltantes": secciones_faltantes,
#         "puntuacion_estructura": round(puntuacion, 2),
#         "contenido_secciones": contenido_secciones
#     }


# # hybrid/structure_analyzer.py
# import fitz  # PyMuPDF
# import re

# def discover_sections_programmatic(pdf_path):
#     """
#     Analiza el PDF buscando títulos basados en marcadores internos 
#     o en características visuales (texto en negrita o tamaño mayor).
#     """
#     print(f"\nAnalizando estructura nativa del PDF: {pdf_path}")
#     secciones_detectadas = []
    
#     try:
#         doc = fitz.open(pdf_path)
        
#         # Estrategia 1: Intentar leer los marcadores/TOC nativos del PDF
#         toc = doc.get_toc()
#         if toc:
#             # El TOC devuelve [nivel, titulo, pagina]. Nos interesan los títulos principales (nivel 1)
#             secciones_detectadas = [item[1].strip() for item in toc if item[0] == 1]
#             print("Secciones detectadas mediante Marcadores (TOC) nativos.")
        
#         # Estrategia 2: Si no hay TOC, analizamos los estilos visuales (Negritas / Tamaño)
#         if not secciones_detectadas:
#             print(" No se encontró TOC nativo. Analizando fuentes y estilos visuales...")
#             for page in doc:
#                 blocks = page.get_text("dict")["blocks"]
#                 for b in blocks:
#                     if "lines" in b:
#                         for l in b["lines"]:
#                             for s in l["spans"]:
#                                 texto = s["text"].strip()
#                                 # Filtrar textos vacíos o muy cortos (números de página, etc.)
#                                 if len(texto) < 4:
#                                     continue
                                
#                                 # Si el texto está en Mayúsculas y es "Bold" (Negrita), o tiene tamaño grande
#                                 es_negrita = "bold" in s["font"].lower() or "black" in s["font"].lower()
#                                 es_mayuscula = texto.isupper()
                                
#                                 if (es_negrita and es_mayuscula) or "RECA" in texto.upper():
#                                     # Evitar duplicados y limpiar espacios
#                                     if texto not in secciones_detectadas:
#                                         # Si es una línea que pertenece a leyes o anexos conocidos
#                                         secciones_detectadas.append(texto)
                                        
#         doc.close()
#     except Exception as e:
#         print(f"Error en la extracción programática: {e}")
        
#     # Limpieza final: Filtrar ruidos comunes si se colaron párrafos
#     secciones_limpias = [s for s in secciones_detectadas if len(s) < 100]
#     print(f"Secciones identificadas automáticamente: {secciones_limpias}\n")
#     return list(dict.fromkeys(secciones_limpias))  # Elimina duplicados manteniendo orden

# def analyze_document_structure_dynamic(actual_text, secciones_esperadas):
#     """Mantiene la lógica exacta de división y recorte de contenidos intermedios."""
#     secciones_encontradas = []
#     secciones_faltantes = []
#     contenido_secciones = {}
#     matches = []
    
#     for nombre_seccion in secciones_esperadas:
#         if "RECA" in nombre_seccion.upper():
#             regex = r"RECA:\s*.*"
#         else:
#             regex = re.escape(nombre_seccion)
            
#         match = re.search(regex, actual_text, re.IGNORECASE)
        
#         if match:
#             texto_real_encontrado = match.group(0).strip()
#             matches.append({
#                 "nombre": nombre_seccion,
#                 "inicio": match.start(),
#                 "fin": match.end(),
#                 "texto_titulo": texto_real_encontrado
#             })
#             secciones_encontradas.append(texto_real_encontrado)
#         else:
#             secciones_faltantes.append(nombre_seccion)

#     matches = sorted(matches, key=lambda x: x["inicio"])
    
#     for index, match_actual in enumerate(matches):
#         inicio_contenido = match_actual["fin"]
        
#         if index + 1 < len(matches):
#             fin_contenido = matches[index + 1]["inicio"]
#             texto_interno = actual_text[inicio_contenido:fin_contenido].strip()
#         else:
#             texto_interno = actual_text[inicio_contenido:].strip()
            
#         contenido_secciones[match_actual["nombre"]] = texto_interno

#     total = len(secciones_esperadas)
#     encontradas = len(secciones_encontradas)
#     puntuacion = (encontradas / total) * 100 if total > 0 else 0
    
#     return {
#         "secciones_esperadas": secciones_esperadas,
#         "secciones_encontradas": secciones_encontradas,
#         "secciones_faltantes": secciones_faltantes,
#         "puntuacion_estructura": round(puntuacion, 2),
#         "contenido_secciones": contenido_secciones
#     }

# # hybrid/structure_analyzer.py
# import re

# def discover_sections_programmatic(pdf_path=None):
#     """
#     Devuelve la lista maestra de leyes y códigos regulatorios esperados
#     en el Anexo de Disposiciones Legales.
#     """
#     # Lista de control fija basada en los documentos corporativos estándar
#     return [
#         "Ley General de Organizaciones y Actividades Auxiliares del Crédito",
#         "Ley General de Títulos y Operaciones de Crédito",
#         "Ley para la Transparencia y Ordenamiento de los Servicios Financieros",
#         "Ley para Regular las Sociedades de Información Crediticia",
#         "Ley de Instituciones de Crédito",
#         "Ley de Protección y Defensa al Usuario de Servicios Financieros",
#         "Código Civil Federal"
#     ]

# def analyze_document_structure_dynamic(actual_text, secciones_esperadas):
#     """
#     Segmenta el texto del PDF dinámicamente buscando la posición de cada ley.
#     """
#     secciones_encontradas = []
#     secciones_faltantes = []
#     contenido_secciones = {}
#     matches = []
    
#     # Buscar cada ley en el texto
#     for nombre_seccion in secciones_esperadas:
#         # Normalizar espacios para evitar problemas de saltos de línea
#         regex = re.escape(nombre_seccion)
#         match = re.search(regex, actual_text, re.IGNORECASE)
        
#         if match:
#             matches.append({
#                 "nombre": nombre_seccion,
#                 "inicio": match.start(),
#                 "fin": match.end()
#             })
#             secciones_encontradas.append(nombre_seccion)
#         else:
#             secciones_faltantes.append(nombre_seccion)

#     # Ordenar los cortes según su aparición en el documento
#     matches = sorted(matches, key=lambda x: x["inicio"])
    
#     # Extraer el texto que queda entre una ley y la siguiente
#     for index, match_actual in enumerate(matches):
#         inicio_contenido = match_actual["fin"]
        
#         if index + 1 < len(matches):
#             fin_contenido = matches[index + 1]["inicio"]
#             texto_interno = actual_text[inicio_contenido:fin_contenido].strip()
#         else:
#             texto_interno = actual_text[inicio_contenido:].strip()
            
#         contenido_secciones[match_actual["nombre"]] = texto_interno

#     total = len(secciones_esperadas)
#     encontradas = len(secciones_encontradas)
#     puntuacion = (encontradas / total) * 100 if total > 0 else 0
    
#     return {
#         "secciones_esperadas": secciones_esperadas,
#         "secciones_encontradas": secciones_encontradas,
#         "secciones_faltantes": secciones_faltantes,
#         "puntuacion_estructura": round(puntuacion, 2),
#         "contenido_secciones": contenido_secciones
#     }



# # hybrid/structure_analyzer.py
# import fitz  # PyMuPDF
# import re

# def discover_sections_programmatic(pdf_path):
#     """
#     Detecta los títulos del documento analizando las propiedades visuales reales:
#     Texto en NEGRITA (Bold) y posicionado de forma CENTRADA en la página.
#     """
#     print(f"\nAnalizando diseño visual del PDF para detectar títulos: {pdf_path}")
#     secciones_detectadas = []
    
#     try:
#         doc = fitz.open(pdf_path)
        
#         for page in doc:
#             # Obtener el ancho de la página para calcular el centro matemático
#             ancho_pagina = page.rect.width
#             centro_pagina = ancho_pagina / 2
            
#             # Extraemos el texto estructurado con detalles de fuentes y posiciones
#             blocks = page.get_text("dict")["blocks"]
            
#             for b in blocks:
#                 if "lines" in b:
#                     for l in b["lines"]:
#                         for s in l["spans"]:
#                             texto = s["text"].strip()
                            
#                             # Filtro inicial: omitir ruidos, números sueltos o textos muy cortos
#                             if len(texto) < 6 or len(texto) > 100:
#                                 continue
                                
#                             # 1. Validación de NEGRITA (Buscamos 'bold' o 'black' en el nombre de la fuente)
#                             fuente = s["font"].lower()
#                             es_negrita = "bold" in fuente or "black" in fuente
                            
#                             # 2. Validación de CENTRADO matemático
#                             # x0 es donde inicia el texto a la izquierda, x1 donde termina a la derecha
#                             x0, x1 = s["bbox"][0], s["bbox"][2]
#                             centro_texto = (x0 + x1) / 2
                            
#                             # Tolerancia de pixeles para el centrado (margen de error aceptable de 15px)
#                             es_centrado = abs(centro_texto - centro_pagina) < 15.0
                            
#                             # 3. Filtro extra: Omitir si es una URL o marcas del header/footer conocido
#                             if "http" in texto or "RECA" in texto or "Estructuras" in texto:
#                                 continue

#                             if es_negrita and es_centrado:
#                                 if texto not in secciones_detectadas:
#                                     secciones_detectadas.append(texto)
                                    
#         doc.close()
#     except Exception as e:
#         print(f"Error en la inspección visual de fuentes: {e}")
        
#     print(f"Secciones visuales (Negrita + Centrado) detectadas: {secciones_detectadas}\n")
#     return secciones_detectadas


# def analyze_document_structure_dynamic(actual_text, secciones_esperadas):
#     """
#     Segmenta el contenido usando los títulos visuales detectados.
#     """
#     secciones_encontradas = []
#     secciones_faltantes = []
#     contenido_secciones = {}
#     matches = []
    
#     for nombre_seccion in secciones_esperadas:
#         # Buscamos la coincidencia exacta del título de la sección en el texto plano
#         regex = re.escape(nombre_seccion)
#         match = re.search(regex, actual_text, re.IGNORECASE)
        
#         if match:
#             matches.append({
#                 "nombre": nombre_seccion,
#                 "inicio": match.start(),
#                 "fin": match.end()
#             })
#             secciones_encontradas.append(nombre_seccion)
#         else:
#             secciones_faltantes.append(nombre_seccion)

#     # Ordenamos cronológicamente según aparecen en el documento
#     matches = sorted(matches, key=lambda x: x["inicio"])
    
#     for index, match_actual in enumerate(matches):
#         inicio_contenido = match_actual["fin"]
        
#         if index + 1 < len(matches):
#             fin_contenido = matches[index + 1]["inicio"]
#             texto_interno = actual_text[inicio_contenido:fin_contenido].strip()
#         else:
#             texto_interno = actual_text[inicio_contenido:].strip()
            
#         contenido_secciones[match_actual["nombre"]] = texto_interno

#     total = len(secciones_esperadas)
#     encontradas = len(secciones_encontradas)
#     puntuacion = (encontradas / total) * 100 if total > 0 else 0
    
#     return {
#         "secciones_esperadas": secciones_esperadas,
#         "secciones_encontradas": secciones_encontradas,
#         "secciones_faltantes": secciones_faltantes,
#         "puntuacion_estructura": round(puntuacion, 2),
#         "contenido_secciones": contenido_secciones
#     }

# # hybrid/structure_analyzer.py
# import os
# import fitz
# from docx import Document

# def _discover_sections_from_pdf(pdf_path):
#     """Lógica original basada en PyMuPDF para archivos PDF."""
#     secciones = []
#     try:
#         doc = fitz.open(pdf_path)
#         for page in doc:
#             # Puedes mantener aquí tu lógica actual basada en fuentes o coordenadas
#             # Por ejemplo, una búsqueda rápida de líneas cortas en mayúsculas/negritas:
#             text_instances = page.get_text("blocks")
#             for block in text_instances:
#                 linea = block[4].strip()
#                 # Filtro genérico descriptivo (ej: PALABRAS EN MAYÚSCULAS)
#                 if linea.isupper() and len(linea) < 50 and not linea.isdigit():
#                     secciones.append(linea)
#         doc.close()
#     except Exception:
#         pass
#     return list(set(secciones))

# def _discover_sections_from_docx(docx_path):
#     """
#     Lógica nativa para detectar títulos/secciones en archivos de Word
#     evaluando si el párrafo está centrado o formateado en negrita (Bold).
#     """
#     secciones = []
#     try:
#         doc = Document(docx_path)
#         for paragraph in doc.paragraphs:
#             texto = paragraph.text.strip()
#             if not texto:
#                 continue
                
#             # 1. Validar si el párrafo completo está marcado con estilo de título
#             # o si su alineación es CENTRADA (WD_ALIGN_PARAGRAPH.CENTER suele ser 1)
#             is_centered = paragraph.alignment == 1 or (paragraph.style and "heading" in paragraph.style.name.lower())
            
#             # 2. Validar si los fragmentos (runs) del texto vienen en negrita
#             is_bold = any(run.bold for run in paragraph.runs if run.text.strip())
            
#             # Si cumple las condiciones de estructura de título, lo guardamos
#             if (is_bold or is_centered) and len(texto) < 60:
#                 secciones.append(texto)
#     except Exception:
#         pass
#     return secciones

# def discover_sections_programmatic(file_path):
#     """
#     DETECTOR ESTRUCTURAL DINÁMICO: Identifica el tipo de archivo 
#     y aplica las reglas de estilo correctas para extraer encabezados/secciones.
#     """
#     ext = os.path.splitext(file_path)[1].lower()
    
#     if ext == '.pdf':
#         return _discover_sections_from_pdf(file_path)
#     elif ext == '.docx':
#         return _discover_sections_from_docx(file_path)
#     return []

# # hybrid/structure_analyzer.py (Sección interna de docx)
# import os
# from docx import Document

# def _discover_sections_from_docx(docx_path):
#     """
#     Detector flexible de secciones en Word. Captura títulos basados en mayúsculas,
#     estilos de encabezado o marcas de negrita.
#     """
#     secciones = []
#     try:
#         doc = Document(docx_path)
#         for paragraph in doc.paragraphs:
#             texto = paragraph.text.strip()
#             if not texto or len(texto) > 60:
#                 continue
                
#             # Criterios flexibles de títulos:
#             is_uppercase = texto.isupper() and not texto.isdigit()
#             is_bold = any(run.bold for run in paragraph.runs if run.text.strip())
#             is_heading_style = paragraph.style and "heading" in paragraph.style.name.lower()
#             is_centered = paragraph.alignment == 1

#             if is_uppercase or is_bold or is_heading_style or is_centered:
#                 if texto not in secciones:
#                     secciones.append(texto)
#     except Exception:
#         pass
#     return secciones

# def _discover_sections_from_pdf(pdf_path):
#     # (Mantén aquí tu lógica original de PDF que ya te funcionaba bien)
#     import fitz
#     secciones = []
#     try:
#         doc = fitz.open(pdf_path)
#         for page in doc:
#             text_instances = page.get_text("blocks")
#             for block in text_instances:
#                 linea = block[4].strip()
#                 if linea.isupper() and len(linea) < 50 and not linea.isdigit():
#                     secciones.append(linea)
#         doc.close()
#     except Exception:
#         pass
#     return list(set(secciones))

# def discover_sections_programmatic(file_path):
#     ext = os.path.splitext(file_path)[1].lower()
#     if ext == '.pdf':
#         return _discover_sections_from_pdf(file_path)
#     elif ext == '.docx':
#         return _discover_sections_from_docx(file_path)
#     return []


# hybrid/structure_analyzer.py
import os
import fitz
from docx import Document

def _discover_sections_from_pdf(pdf_path):
    """Lógica optimizada para PDF: Detecta títulos reales ignorando celdas de tablas o fechas."""
    secciones = []
    try:
        doc = fitz.open(pdf_path)
        for page in doc:
            text_instances = page.get_text("blocks")
            for block in text_instances:
                linea = block[4].strip()
                
                # FILTROS DE LIMPIEZA CRÍTICOS:
                if not linea:
                    continue
                if len(linea) > 50:
                    continue
                # Si contiene signos de dinero, variables de fecha o diagonales, NO es un título
                if "$" in linea or "/" in linea or "dd/mm" in linea.lower() or "pago no" in linea.lower():
                    continue
                # Si es puramente un número suelto
                if linea.isdigit():
                    continue
                    
                # Si pasa los filtros y es mayúscula pura o contiene las etiquetas de cierre
                if (linea.isupper() and not linea.isdigit()) or "“EL " in linea:
                    # Limpiar saltos de línea internos que rompen el JSON
                    linea_limpia = " ".join(linea.split())
                    if linea_limpia not in secciones:
                        secciones.append(linea_limpia)
        doc.close()
    except Exception:
        pass
    return secciones

def _discover_sections_from_docx(docx_path):
    """Detector flexible de secciones en Word."""
    secciones = []
    try:
        doc = Document(docx_path)
        for paragraph in doc.paragraphs:
            texto = paragraph.text.strip()
            if not texto or len(texto) > 50:
                continue
                
            is_uppercase = texto.isupper() and not texto.isdigit()
            is_bold = any(run.bold for run in paragraph.runs if run.text.strip())
            is_heading_style = paragraph.style and "heading" in paragraph.style.name.lower()
            is_centered = paragraph.alignment == 1

            if is_uppercase or is_bold or is_heading_style or is_centered:
                if "$" in texto or "/" in texto:
                    continue
                if texto not in secciones:
                    secciones.append(texto)
    except Exception:
        pass
    return secciones

def discover_sections_programmatic(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    if ext == '.pdf':
        return _discover_sections_from_pdf(file_path)
    elif ext == '.docx':
        return _discover_sections_from_docx(file_path)
    return []